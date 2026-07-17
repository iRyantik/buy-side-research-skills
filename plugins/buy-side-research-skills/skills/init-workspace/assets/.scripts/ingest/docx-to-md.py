#!/usr/bin/env python3
"""docx-to-md.py — Convert Word document to markdown.

Usage:
    python docx-to-md.py <input.docx> [--output <path>]
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path


def convert(input_path: Path, output_path: Path) -> dict:
    from docx import Document

    doc = Document(str(input_path))
    parts = [f"# {input_path.name}", ""]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            parts.append("")
            continue
        # Preserve heading levels
        if para.style.name.startswith("Heading"):
            level = para.style.name.split()[-1]
            try:
                level = int(level)
            except ValueError:
                level = 2
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)
        parts.append("")

    for idx, table in enumerate(doc.tables, 1):
        parts.append(f"## Table {idx}")
        parts.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
            if i == 0:
                parts.append("| " + " | ".join("---" for _ in cells) + " |")
        parts.append("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    text = "\n".join(parts)
    output_path.write_text(text, encoding="utf-8")

    return {
        "status": "converted",
        "engine": "python-docx",
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "chars": len(text),
    }


def main():
    parser = argparse.ArgumentParser(description="Convert Word document to markdown")
    parser.add_argument("input", help="DOCX file path")
    parser.add_argument("--output", "-o", help="Output markdown path")
    args = parser.parse_args()

    src = Path(args.input)
    if not src.exists():
        print(f"ERROR: file not found: {src}", file=sys.stderr)
        sys.exit(1)

    dst = Path(args.output) if args.output else src.with_suffix(".md")
    result = convert(src, dst)
    print(f"OK {result['paragraphs']} paras, {result['tables']} tables, {result['chars']:,} chars → {dst}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
