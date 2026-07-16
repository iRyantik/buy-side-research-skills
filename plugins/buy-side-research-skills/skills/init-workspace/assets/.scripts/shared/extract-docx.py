#!/usr/bin/env python3
"""extract-docx — Word document to text/markdown. Lightweight.

Usage:
  python extract-docx.py <file>              # plain text
  python extract-docx.py <file> --markdown    # markdown with headings/tables
  python extract-docx.py <file> --json        # structured JSON
"""
from __future__ import annotations

import argparse
import json
import re
import sys


def _clean(text: str) -> str:
    return text.strip()


def extract(filepath: str) -> dict:
    """Extract paragraphs and tables. Returns {paragraphs, tables}."""
    try:
        import docx
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install --user python-docx", file=sys.stderr)
        sys.exit(1)

    doc = docx.Document(filepath)
    paragraphs = []
    tables = []

    # Extract paragraphs with style info
    for para in doc.paragraphs:
        text = _clean(para.text)
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        is_heading = style.startswith("Heading") or style.startswith("heading")
        level = 0
        if is_heading:
            m = re.search(r'(\d)', style)
            level = int(m.group(1)) if m else 1
        paragraphs.append({
            "text": text,
            "style": style,
            "heading": is_heading,
            "level": level,
        })

    # Extract tables
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [_clean(cell.text) for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append({"index": i, "headers": rows[0] if rows else [], "rows": rows[1:] if len(rows) > 1 else []})

    return {"paragraphs": paragraphs, "tables": tables}


def to_markdown(data: dict) -> str:
    """Convert extracted content to markdown."""
    lines = []
    for p in data["paragraphs"]:
        if p["heading"]:
            prefix = "#" * p["level"]
            lines.append(f"{prefix} {p['text']}")
            lines.append("")
        else:
            lines.append(p["text"])
            lines.append("")

    for t in data["tables"]:
        if t["headers"]:
            lines.append("| " + " | ".join(t["headers"]) + " |")
            lines.append("|---" * len(t["headers"]) + "|")
            for row in t["rows"]:
                lines.append("| " + " | ".join(row) + " |")
        lines.append("")

    return "\n".join(lines)


def main():
    p = argparse.ArgumentParser(description="DOCX to text/markdown/JSON")
    p.add_argument("file", help="DOCX file path")
    p.add_argument("--markdown", action="store_true", help="Output as markdown")
    p.add_argument("--json", action="store_true", help="Output as structured JSON")
    args = p.parse_args()

    data = extract(args.file)

    if args.json:
        print(json.dumps(data, ensure_ascii=False, indent=2))
    elif args.markdown:
        print(to_markdown(data))
    else:
        for p in data["paragraphs"]:
            print(p["text"])


if __name__ == "__main__":
    main()
